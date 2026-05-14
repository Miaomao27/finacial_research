#!/usr/bin/env python3
"""
Template: Generate academic paper .docx with embedded figures + metaso references.
Usage: 
  1. Run analysis and generate charts first
  2. Use metaso to search for academic references
  3. Edit OUTPUT_DIR and content sections below
  4. Run: python generate_paper_docx.py

Requires: pip install python-docx
          metaso MCP access for academic reference search
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ===== CONFIG — EDIT =====
OUTPUT_DIR = '/path/to/your/results/dir'
STOCK_NAME = 'Agricultural Bank (601288.SH)'
DATA_RANGE = '2010-07-15 to 2026-05-13'
N_OBS = 3833

FIGURE_PATHS = [
    ('fig1_daily_boxplot.png', '图1  日收益率按五行分组箱线图'),
    ('fig2_weekly_rolling.png', '图2  五行分组周收益率3年滚动窗口差异'),
    ('fig3_monthly_train_test.png', '图3  月收益率训练集vs测试集对比'),
    ('fig4_summary_all.png', '图4  全部假设检验p值汇总'),
]

# ===== METASO REFERENCE WORKFLOW =====
# Before finalizing, search for references:
"""
from hermes_tools import mcp_metaso_metaso_chat

# Search for Chinese academic literature
refs = mcp_metaso_metaso_chat(
    message="搜索中国股市日历效应、天干地支择时的中英文学术文献",
    model="fast"
)
# Metaso returns citations with authors, date, title, link
# Format them as reference strings and add to the refs list below
"""

# Metaso-sourced references from the Agricultural Bank study:
METASO_REFS = [
    '[1] Cross, F. (1973). The behavior of stock prices on Fridays and Mondays. Financial Analysts Journal, 29(6), 67-69.',
    '[2] 薛继锐, 顾岚. (2000). 中国股票市场的日历效应分析. 数理统计与管理, 19(2), 10-15.',
    '[3] 郑雅芹. (2019). 不同市态下中国股市的日历效应研究 [硕士论文]. 上海外国语大学.',
    '[4] 安信证券. (2021). 天干地支在择时中的应用初探. 证券研究报告.',
    '[5] Kruskal, W.H. & Wallis, W.A. (1952). Use of ranks in one-criterion variance analysis. JASA, 47(260), 583-621.',
    '[6] Newey, W.K. & West, K.D. (1987). A consistent covariance matrix. Econometrica, 55(3), 703-708.',
    '[7] Fama, E.F. (1970). Efficient capital markets. The Journal of Finance, 25(2), 383-417.',
    '[8] Lo, A.W. & MacKinlay, A.C. (1999). A Non-Random Walk Down Wall Street. Princeton UP.',
]

# ===== BUILD DOCUMENT =====
doc = Document()

style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)

# Title
title = doc.add_heading('论文标题', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('副标题')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

# Abstract
doc.add_heading('摘  要', level=1)
doc.add_paragraph(
    f'本研究以{STOCK_NAME}上市的{DATA_RANGE}共{N_OBS}个交易日的日线数据为样本，'
    f'构建了日、周、月三层频率的统计检验框架，系统检验假设的预测能力。'
    '研究采用严格的样本外验证、混杂因素剥离和稳健性检验。结论：无显著预测能力。'
)

# 1 Introduction
doc.add_heading('1  引言', level=1)
doc.add_paragraph(
    'Include motivation: why test this hypothesis, existing literature, research gap. '
    'For the 天干地支案例: mention the 2021 安信证券 controversial report as the direct motivation.'
)

# 2 Data & Method
doc.add_heading('2  数据与方法', level=1)
doc.add_paragraph(f'Data: {STOCK_NAME}, {DATA_RANGE}, {N_OBS} observations.')
doc.add_paragraph(
    'Three-layer frequency framework: daily (baseline), weekly (rolling window), '
    'monthly (train/test split). '
    'OLS regression with Newey-West standard errors for confounding detection.'
)

# 3 Results
doc.add_heading('3  结果', level=1)

sections = [
    ('3.1  日频分析', 0, 'Daily results text...'),
    ('3.2  周频分析', 1, 'Weekly rolling window results...'),
    ('3.3  月频分析', 2, 'Monthly train/test split results...'),
]

for heading, fig_idx, text in sections:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(text)
    img_path = os.path.join(OUTPUT_DIR, FIGURE_PATHS[fig_idx][0])
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph(FIGURE_PATHS[fig_idx][1])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)

doc.add_heading('3.4  混杂效应剥离', level=2)
doc.add_paragraph('Nested regression results: M1 (category only) → M2 (+calendar month) → M3 (+year).')

# Regression table
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'R²', 'F-test p', 'Category vars sig?']):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate([
    ['Only category', '0.03', '0.206', '—'],
    ['+ Calendar month', '0.10', '0.038', 'No'],
    ['+ Year', '0.21', '<0.001', 'No'],
]):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_heading('3.5  稳健性检验', level=2)
doc.add_paragraph('Winsorization results...')

img_path = os.path.join(OUTPUT_DIR, FIGURE_PATHS[3][0])
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph(FIGURE_PATHS[3][1])
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)

# 4 Discussion
doc.add_heading('4  讨论', level=1)
doc.add_paragraph(
    'Why the hypothesis fails: calendar effect confounding. '
    'For 天干地支: different 五行 map to different months by construction. '
    'The "pattern" is just seasonal A-share effects wearing a traditional label.'
)

# 5 Conclusion
doc.add_heading('5  结论', level=1)
doc.add_paragraph('Final verdict: null result — no predictive power detected.')

# References
doc.add_heading('参考文献', level=1)
for ref in METASO_REFS:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

output_path = os.path.join(OUTPUT_DIR, '学术论文.docx')
doc.save(output_path)
print(f'Paper saved to: {output_path}')
print(f'Included {len(FIGURE_PATHS)} figures, {len(METASO_REFS)} references (metaso-sourced)')
