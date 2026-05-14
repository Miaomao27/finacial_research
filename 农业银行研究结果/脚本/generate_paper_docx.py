#!/usr/bin/env python3
"""
生成完整的中文学术论文 .docx
嵌入4张图表 + 结果表格 + 10篇参考文献
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = '/home/cpy/文档/金融数据库建立/研究结果'
CHART_DIR = os.path.join(OUTPUT_DIR, '图表')

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# Heading styles
for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(3)

def add_paragraph(text, bold=False, size=None, align=None, indent=True, font_name=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_centered_text(text, size=12, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading)

def format_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
#  标题页
# ============================================================
doc.add_paragraph()  # spacing
doc.add_paragraph()
doc.add_paragraph()

add_centered_text('五行炒股指南？', size=26, bold=True)
add_centered_text('——农业银行股价与天干地支五行关系的统计学检验', size=16, bold=False,
                   color=RGBColor(80, 80, 80))

doc.add_paragraph()
doc.add_paragraph()

add_centered_text('Agentic Financial Research Laboratory', size=12)
add_centered_text('2026年5月', size=12, color=RGBColor(128, 128, 128))

doc.add_page_break()

# ============================================================
#  摘要
# ============================================================
doc.add_heading('摘  要', level=1)

add_paragraph(
    '天干地支五行理论在中国传统文化中影响深远，近年来甚至有投资者将其用于A股择时。'
    '本文以农业银行（601288.SH）2010年7月至2026年5月共计3,833个交易日的日线数据为样本，'
    '构建了日、周、月三层频率的递进统计检验框架，系统检验五行属性是否对银行股价格具有预测能力。'
    '研究采用Kruskal-Wallis秩和检验、Newey-West稳健标准误回归、'
    '训练集/测试集分割（70%/30%）以及缩尾处理等多种方法，从多角度验证假设。'
)

add_paragraph(
    '研究发现：全部9个假设检验在5%显著性水平下均不显著；'
    '训练集上最强的"规律"（土组收益最高、火组最低）在测试集中方向完全反转，确认为过拟合产物；'
    '回归分析表明五行变量在加入阳历月份控制后即丧失解释力，'
    '说明五行本质上是历法季节性的代理变量。'
    '结论：天干地支五行对农业银行股价不具备稳健的预测能力，'
    '所谓"五行择时"更可能是日历效应与多重比较偏误的叠加结果。'
)

add_paragraph(
    '关键词：天干地支；五行理论；日历效应；Kruskal-Wallis检验；Newey-West回归；样本外验证',
    bold=True, indent=False
)

doc.add_page_break()

# ============================================================
#  1 引言
# ============================================================
doc.add_heading('1  引言', level=1)

add_paragraph(
    '在中国金融市场的民间智慧中，天干地支择时始终占有一席之地。'
    '从"甲木生火，火克金"的传统辩证逻辑，到社交媒体上"金日买银行，水日买科技"的经验之谈，'
    '五行学说被赋予了超越其历法本源的金融预测功能。'
    '2021年某券商发布题为《天干地支在择时中的应用初探》的争议研报[7]，'
    '试图用量化模型验证"五行与A股走势的关系"，一度引发市场热议和监管关注。'
    '该研报的核心逻辑——将天干地支的五行属性与股票收益率建立直接映射——'
    '在学术层面上从未经过严格的统计检验。'
)

add_paragraph(
    '然而，传统智慧与统计显著性之间往往存在巨大鸿沟。'
    '天干地支本质上是中国传统历法系统的组成部分，与阳历月份存在周期性对应关系。'
    '如果五行"规律"在剥离日历效应后即告消失，那么它不过是季节性的华丽包装，'
    '不具备真正的预测价值。这一辨伪逻辑在金融研究中并非新事——'
    '从20世纪70年代Cross[1]对"周末效应"的开创性分析，'
    '到Fama[5]对有效市场假说的系统论述，'
    '再到Lo与MacKinlay[9]对"非随机游走"的深入剖析，'
    '市场异象的样本外失效始终是金融研究的核心议题。'
)

add_paragraph(
    '本研究旨在通过严格的统计方法回答一个简单而根本的问题：'
    '天干地支五行的信息能否转化为统计上显著的超额收益？'
    '选择农业银行（601288.SH）作为研究对象出于以下考虑：'
    '（1）银行股波动相对平稳，不易受个股噪声干扰；'
    '（2）上市时间长（2010年至今），样本量充足；'
    '（3）日均换手率低，更有利于检测结构性而非流动性的效应。'
)

# ============================================================
#  2 数据与方法
# ============================================================
doc.add_heading('2  数据与方法', level=1)

doc.add_heading('2.1  数据来源', level=2)

add_paragraph(
    '本研究使用MySQL数据库（china_finance_db）存储的农业银行（601288.SH）日线数据，'
    '包含3,833条真实交易记录，时间跨度为2010年7月15日至2026年5月13日（约16年）。'
    '行情数据涵盖开盘价、收盘价、最高价、最低价、成交量及成交额，'
    '复权采用后复权方式以保持收益率计算的一致性。'
)

add_paragraph(
    '每个交易日的干支标注通过lunar_python天文历法库精确计算，包括年柱、月柱、日柱的天干和地支。'
    '五行映射规则严格按照传统五行分类：'
    '天干方面，甲乙→木、丙丁→火、戊己→土、庚辛→金、壬癸→水；'
    '地支方面，寅卯→木、巳午→火、申酉→金、亥子→水、辰戌丑未→土。'
    '月柱的划分按节气进行：寅月（立春至惊蛰）、卯月（惊蛰至清明）、辰月（清明至立夏），以此类推。'
)

add_paragraph(
    '在日线数据基础上，进一步聚合生成周频数据（811条）和月频数据（191条）。'
    '周收益率按每周最后一个交易日的收盘价相对于前一周最后一个交易日计算；'
    '月收益率同理。三个时间尺度构成递进的分析框架，'
    '确保结论在不同频率上的一致性。'
)

doc.add_heading('2.2  分析方法', level=2)

add_paragraph(
    '研究采用五层递进的分析策略，力求从多个角度验证或否定五行与股价的关联：'
)

# Bullet points for methods
methods = [
    ('日频基线检验', '以全部3,833个交易日为样本，按天干五行和地支五行分别分组，'
     '使用Kruskal-Wallis秩和检验[3]比较5个组的收益率分布是否存在显著差异。'
     '该检验为非参数方法，不假设正态分布，适用于金融收益率数据。'),
    ('周频滚动窗口分析', '在811条周度数据上重复K-W检验，同时引入3年（约156周）滚动窗口，'
     '计算各五行对组合的收益率差异时序，观察差异方向是否稳定持久。'),
    ('月频样本外验证', '将191个月度数据按时序分割为70%训练集（2010年7月至2021年7月，'
     '133个月）和30%测试集（2021年8月至2026年5月，58个月）。'
     '在训练集上充分探索各种可能的五行规律，固定一个最强假设后，'
     '原封不动地在测试集上检验其预测能力。这是本研究的核心辨伪机制。'),
    ('混杂效应剥离', '建立OLS回归模型，以月收益率为因变量，'
     '天干五行虚拟变量为核心解释变量，逐次加入阳历月份和年份虚拟变量作为控制变量。'
     '标准误采用Newey-West[4]方法进行异方差和自相关一致性修正。'
     '若五行变量在加入月份控制后丧失显著性，则表明五行只是日历效应的代理变量。'),
    ('稳健性检验', '对日收益率做1%和99%分位数缩尾（winsorize）处理后，'
     '重新执行日频天干五行K-W检验，确认结论对极端值是否敏感。'),
]

for title, desc in methods:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(f'（{methods.index((title, desc)) + 1}）{title}：')
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.line_spacing = 1.5

# ============================================================
#  3 结果
# ============================================================
doc.add_heading('3  结果', level=1)

# ---- 3.1 日频分析 ----
doc.add_heading('3.1  日频分析——五行分组收益率几乎完全重合', level=2)

add_paragraph(
    '表1汇总了日收益率按天干五行分组的基本统计量。'
    '5个五行组的日均收益率高度集中在+0.02%至+0.07%之间，'
    '标准差均在1.17%至1.32%的狭窄区间内，各组上涨天数占比约40%至44%，'
    '组间差异极小。Kruskal-Wallis检验统计量H=0.95（p=0.917），'
    '差异完全不显著。p=0.917是所有统计检验可获得的最高p值之一——'
    '在0到1的区间里，这是统计学能给出的"没有关系"的最强信号。'
)

# 表1: 天干五行日收益率统计
add_paragraph('表1  日收益率按天干五行分组统计', bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

table1 = doc.add_table(rows=7, cols=6)
table1.style = 'Light Grid Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['五行', '样本数', '日均收益率(%)', '标准差(%)', '中位数(%)', '上涨占比(%)']
for i, h in enumerate(headers1):
    format_cell(table1.rows[0].cells[i], h, bold=True, size=9)
    set_cell_shading(table1.rows[0].cells[i], 'D9E2F3')

data1 = [
    ['金', '768', '+0.069', '1.249', '0.000', '43.5'],
    ['木', '770', '+0.040', '1.325', '0.000', '42.7'],
    ['水', '760', '+0.054', '1.271', '0.000', '44.2'],
    ['火', '770', '+0.046', '1.267', '0.000', '40.1'],
    ['土', '765', '+0.019', '1.165', '0.000', '44.4'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        format_cell(table1.rows[r+1].cells[c], val, size=9)

# K-W row
format_cell(table1.rows[6].cells[0], 'K-W检验', bold=True, size=9)
table1.rows[6].cells[1].merge(table1.rows[6].cells[3])
format_cell(table1.rows[6].cells[1], 'H=0.95, p=0.917', bold=True, size=9)
table1.rows[6].cells[4].merge(table1.rows[6].cells[5])
format_cell(table1.rows[6].cells[4], '不显著 (α=0.05)', size=9)
set_cell_shading(table1.rows[6].cells[0], 'E2EFDA')

add_paragraph('', indent=False)  # spacer

add_paragraph(
    '地支五行分组的结果与之类似。Kruskal-Wallis检验H=6.82（p=0.146），'
    '同样不显著。各组收益率的中位数均为零，分布高度重合。'
    '图1的箱线图直观地展示了这一特征——5个箱体几乎完全重叠，'
    '任何肉眼可辨的差异都不足以支撑统计推断。'
)

# 插入图1
doc.add_picture(os.path.join(CHART_DIR, 'fig1_daily_boxplot.png'), width=Inches(5.5))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图1  日收益率按五行分组箱线图。各组中位数均靠近零线，分布高度重合。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# ---- 3.2 周频分析 ----
doc.add_heading('3.2  周频分析——滚动窗口的正负交替', level=2)

add_paragraph(
    '周频分析的Kruskal-Wallis检验同样未能达到显著性阈值：'
    '天干组H=5.83（p=0.212），地支组H=4.81（p=0.307）。'
    '时间聚合并未像某些研究所期待的那样"放大"五行效应，'
    '反而进一步确认了日频层面的结论。'
)

add_paragraph(
    '更具说明力的是3年滚动窗口的分析。'
    '表2列出了各五行对组合的周收益率平均差异。'
    '以差异幅度最大的"火-金"组合为例，其全样本均值约为-36个基点（bps），'
    '意味着火的周收益率平均比金低0.36个百分点——'
    '这似乎是某种意义上的"规律"。'
    '然而，图2的时序图显示，这一差异在历史长河中正负交替，'
    '不存在长期稳定的方向性偏移。'
    '即使在差异最极端的时间段，方向性偏移也无法持续超过一年。'
    '这种围绕零轴随机波动的模式，是典型的噪声而非信号。'
)

# 表2: 周频滚动窗口五行对差异
add_paragraph('表2  五行对周收益率平均差异（全样本，单位：bps）', bold=True, indent=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['五行对', '平均差异(bps)', '出现频次']
for i, h in enumerate(headers2):
    format_cell(table2.rows[0].cells[i], h, bold=True, size=9)
    set_cell_shading(table2.rows[0].cells[i], 'D9E2F3')

data2 = [
    ['火-金', '-36.3', '656周'],
    ['木-金', '-4.3', '656周'],
    ['土-金', '-3.2', '656周'],
    ['木-火', '+32.0', '656周'],
    ['土-火', '+33.1', '656周'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        format_cell(table2.rows[r+1].cells[c], val, size=9)

add_paragraph('', indent=False)

# 插入图2
doc.add_picture(os.path.join(CHART_DIR, 'fig2_weekly_rolling.png'), width=Inches(5.5))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图2  五行分组周收益率3年滚动窗口差异时序。红色水平线为全样本均值，均值的正负方向不稳定。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# ---- 3.3 月频分析 ----
doc.add_heading('3.3  月频分析——样本外验证的辨伪价值', level=2)

add_paragraph(
    '月频分析是本研究的"主战场"。低频数据通常噪声较少，'
    '如果五行真对股价有结构性影响，在月频层面应该体现得最为清晰。'
    '然而，检验结果再次否定了这一预期。'
)

add_paragraph(
    '训练集（2010年7月至2021年7月，133个月）上，天干五行组间差异的K-W检验H=5.31（p=0.257），'
    '地支五行组H=4.79（p=0.310），均不显著。'
    '在训练集内，表现最优的"土"组（月均收益率+1.13%）与最差的"火"组（月均收益率-1.47%）'
    '构成了一个看似合理的初步假设——"土生金，银行属土，土组应该表现好"。'
)

add_paragraph(
    '然而，在测试集（2021年8月至2026年5月，58个月）上，这一"规律"完全反转：'
    '"土"组变为亏损（月均-1.17%），"火"组反而盈利（月均+0.92%）。'
    '这是经典的过拟合案例。测试集K-W检验天干组H=6.90（p=0.141），'
    '地支组H=5.15（p=0.272），同样不显著。'
)

add_paragraph(
    '图3按月收益率展示了天干×地支12个组合的训练集和测试集对比。'
    '图中大量标注为"反转"的组合说明：12个组合中有7个的训练集和测试集收益方向相反，'
    '进一步印证了样本内结果的脆弱性。'
    '如果基于训练集上的发现构建交易策略并在测试集上实施，结果是亏损而非盈利。'
)

# 表3: 月频检验结果汇总
add_paragraph('表3  月频Kruskal-Wallis检验结果汇总', bold=True, indent=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

table3 = doc.add_table(rows=5, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['检验', 'H统计量', 'p值', '显著性(α=0.05)']
for i, h in enumerate(headers3):
    format_cell(table3.rows[0].cells[i], h, bold=True, size=9)
    set_cell_shading(table3.rows[0].cells[i], 'D9E2F3')

data3 = [
    ['月天干（训练集）', '5.31', '0.257', '不显著'],
    ['月地支（训练集）', '4.79', '0.310', '不显著'],
    ['月天干（测试集）', '6.90', '0.141', '不显著'],
    ['月地支（测试集）', '5.15', '0.272', '不显著'],
]
for r, row_data in enumerate(data3):
    for c, val in enumerate(row_data):
        format_cell(table3.rows[r+1].cells[c], val, size=9)

add_paragraph('', indent=False)

# 插入图3
doc.add_picture(os.path.join(CHART_DIR, 'fig3_monthly_train_test.png'), width=Inches(5.5))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图3  月收益率训练集vs测试集对比。红色"反转"标记表示方向不一致的组合。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# ---- 3.4 混杂效应剥离 ----
doc.add_heading('3.4  混杂效应剥离——五行只是日历效应的"华丽外衣"', level=2)

add_paragraph(
    '为检验五行是否只是阳历月份的代理变量，构建了三组OLS回归模型，'
    '标准误采用Newey-West[4]方法进行异方差和自相关一致性修正。'
    '结果见表4。'
)

# 表4: 回归结果
add_paragraph('表4  Newey-West回归结果：逐次加入控制变量', bold=True, indent=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

table4 = doc.add_table(rows=4, cols=4)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

headers4 = ['模型设定', 'R²', 'F检验p值', '五行变量显著性']
for i, h in enumerate(headers4):
    format_cell(table4.rows[0].cells[i], h, bold=True, size=9)
    set_cell_shading(table4.rows[0].cells[i], 'D9E2F3')

data4 = [
    ['仅天干五行虚拟变量', '0.031', '0.206', '—（整体不显著）'],
    ['+ 阳历月份控制', '0.095', '0.038', '全部不显著'],
    ['+ 年份固定效应', '0.210', '<0.001', '全部不显著'],
]
for r, row_data in enumerate(data4):
    for c, val in enumerate(row_data):
        format_cell(table4.rows[r+1].cells[c], val, size=9)

add_paragraph('', indent=False)

add_paragraph(
    '结果显示：仅包含五行变量时模型整体不显著（F检验p=0.206，R²仅0.031），'
    '说明五行变量本身几乎不解释任何收益变化。'
    '加入阳历月份控制后，F检验固然变得显著（p=0.038），'
    '这正是日历效应（即薛继锐、顾岚[2]及郑雅芹[6]所研究的现象）的证据。'
    '但更为关键的是，五行变量的系数全部不显著——'
    '真正驱动收益差异的是阳历月份的季节性，而非五行本身。'
    '加入年份固定效应后R²提升至0.210，'
    '但五行变量仍无显著贡献，说明跨年度的趋势也不能由五行解释。'
)

add_paragraph(
    '这一结果在直觉上也是必然的：天干地支本质上是历法系统，'
    '与阳历月份天然相关。例如，"金"对应的庚辛天干主要出现在秋季（阳历8-9月），'
    '而"火"对应的丙丁天干主要出现在夏季（阳历5-6月）。'
    '在古今月份基本对齐的前提下，任何"显著的五行效应"都不过是不同月份收益差异的重新编码。'
    '这一结论呼应了日历效应文献[1][2][8]的核心发现：'
    '股票收益的季节性主要源于宏观经济周期、政策节奏和投资者行为等真实因素，'
    '而非历法符号本身的玄学属性。'
)

# ---- 3.5 稳健性检验 ----
doc.add_heading('3.5  稳健性检验', level=2)

add_paragraph(
    '对日收益率做1%和99%分位数缩尾处理后重新执行K-W检验，'
    '日天干统计量H=0.95（p=0.917），与原始结果完全一致。'
    '缩尾处理不改变任何定论，说明结论不受极端值的驱动。'
)

add_paragraph(
    '全部9个假设检验的p值汇总如图4。'
    '图中红色虚线为α=0.05显著性阈值，9个条形无一超越该线。'
    '即使在未做任何多重比较校正的宽松条件下，也没有任何一个检验能达到显著。'
    '若采用Harvey、Liu与Zhu[10]建议的更为严格的p值阈值（p<0.003），结论更加稳固。'
)

# 插入图4
doc.add_picture(os.path.join(CHART_DIR, 'fig4_summary_all.png'), width=Inches(5.0))
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图4  全部9个假设检验的p值条形图。红色虚线为α=0.05显著性阈值。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# ============================================================
#  4 讨论
# ============================================================
doc.add_heading('4  讨论', level=1)

add_paragraph(
    '为什么很多投资者"感觉"五行有用？本文认为至少存在三种认知偏误的叠加：'
)

add_paragraph(
    '第一，日历效应混淆。如3.4节所示，天干地支在本质上是历法系统，'
    '不同五行天然对应不同阳历月份。例如，"金"多出现在阳历8-9月的秋季——'
    'A股在政策密集期往往表现较好；"火"多出现在5-6月——'
    '历史上常有利空出尽的季节性调整。当投资者在未控制月份的情况下比较五行收益率时，'
    '他们看到的实际上是不同季节的股市表现差异，而非五行本身的预测力。'
    '这是一种"时间上的幸存者偏差"。'
)

add_paragraph(
    '第二，多重比较偏误。在日、周、月三个频率×天干、地支两个维度×五个五行组的组合空间中，'
    '可供"挖掘"的统计量多达数十个。仅5个组的成对比较就有10种组合，'
    '而本文的滚动窗口框架实质上对数以百计的检验窗口进行了搜索。'
    '在如此庞大的搜索空间中，找到某个"显著"的局部规律几乎是必然的——'
    '即使数据完全是白噪声。本文之所以未发现任何显著结果，恰恰是因为严格遵循了'
    '"训练集探索→测试集验证"的样本外逻辑，而非在全部数据上进行事后解释。'
)

add_paragraph(
    '第三，叙事偏好。人类天生倾向于在随机性中寻找模式[9]。'
    '"土生金，银行属土"的叙事比"这是一组随机噪声"更容易被记住和传播。'
    '当投资者在某个"金日"看到银行股上涨时，他们会记住这个巧合；'
    '而在"水日"银行股同样上涨时，他们并不会意识到这推翻了自己的理论。'
    '这种选择性记忆在缺乏统计框架的直觉判断中几乎不可克服。'
)

add_paragraph(
    '本研究的一项核心学术价值在于：它实证展示了量化金融中屡见不鲜却常被忽略的道理——'
    '如果不做样本外检验，任何历史数据都能"找到"某种模式。'
    '本文的训练集探索阶段挖掘了数十种可能的五行组合规律，'
    '最终固定的"土>火"假设在测试集中灰飞烟灭。'
    '这个教训不仅适用于五行择时，也是所有技术分析和模式识别策略的共同命门。'
    '从Fama[5]的有效市场假说到Lo与MacKinlay[9]的市场异象研究，'
    '金融学术史上已有太多"看起来很美"的策略在样本外验证中被证伪。'
)

add_paragraph(
    '本研究的局限性包括：（1）仅分析了农业银行单只股票，结论向其他行业和个股的推广需要进一步验证；'
    '（2）使用了最基本的收益率指标，未考虑风险调整后的表现（如夏普比率、Alpha等）；'
    '（3）对天干地支的五行映射采用了最主流的分类方法，'
    '其他流派（如纳音五行）的映射规则可能产生不同结果。'
    '然而，考虑到本文在9个检验中均未发现任何接近显著的结果，'
    '五行映射规则的微调不太可能改变总体结论。'
)

# ============================================================
#  5 结论
# ============================================================
doc.add_heading('5  结论', level=1)

add_paragraph(
    '基于农业银行16年、3,833个交易日的完整数据，本研究经过日-周-月三层频率的统计检验、'
    '严格的样本外验证（70/30时间分割）、混杂因素剥离（阳历月份控制）和稳健性检验（缩尾处理），'
    '得出明确结论：天干地支五行对农业银行股价不具有统计上显著的预测能力。'
    '该结论在多个时间尺度和检验方法上高度一致，不因参数选择或样本分割方式而变化。'
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(24)
run = p.add_run('在统计学上，"没有证据"不等于"证据不存在"。')
run.bold = True
p.add_run(
    '但16年的数据、9个严格假设检验、'
    '以及在训练集-测试集框架下确认的过拟合案例，'
    '足以让我们有相当的把握做出上述判断。'
    '天干地支五行理论是中国传统文化中的宝贵遗产，'
    '在哲学、医学、历法等领域具有不可替代的价值；'
    '但它并不是一个有效的股票市场预测工具——至少不是在农业银行这只股票上。'
    '如果一定要在传统智慧与统计证据之间做出选择，数据已经给出了明确的答案。'
)

doc.add_page_break()

# ============================================================
#  参考文献
# ============================================================
doc.add_heading('参考文献', level=1)

refs = [
    '[1] Cross, F. (1973). The behavior of stock prices on Fridays and Mondays. '
    'Financial Analysts Journal, 29(6), 67-69.',
    '[2] 薛继锐, 顾岚. (2000). 中国股票市场的日历效应分析. 数理统计与管理, 19(2), 10-15.',
    '[3] Kruskal, W.H. & Wallis, W.A. (1952). Use of ranks in one-criterion variance analysis. '
    'Journal of the American Statistical Association, 47(260), 583-621.',
    '[4] Newey, W.K. & West, K.D. (1987). A simple, positive semi-definite, '
    'heteroskedasticity and autocorrelation consistent covariance matrix. '
    'Econometrica, 55(3), 703-708.',
    '[5] Fama, E.F. (1970). Efficient capital markets: A review of theory and empirical work. '
    'The Journal of Finance, 25(2), 383-417.',
    '[6] 郑雅芹. (2019). 不同市态下中国股市的日历效应研究——基于EGARCH-M模型实证分析 '
    '[硕士学位论文]. 上海外国语大学.',
    '[7] 安信证券. (2021). 天干地支在择时中的应用初探. 证券研究报告.',
    '[8] 基于日历效应的行业配置策略研究 [硕士学位论文]. 中央民族大学, 2024.',
    '[9] Lo, A.W. & MacKinlay, A.C. (1999). A Non-Random Walk Down Wall Street. '
    'Princeton University Press.',
    '[10] Harvey, C.R., Liu, Y. & Zhu, H. (2016). … and the cross-section of expected returns. '
    'The Review of Financial Studies, 29(1), 5-68.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    if any(c in ref[0:5] for c in ['薛继锐', '郑雅芹', '安信证券', '中央民族大学']):
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
#  保存
# ============================================================
output_path = os.path.join(OUTPUT_DIR, '学术论文_农行天干地支五行检验.docx')
doc.save(output_path)
print(f'论文已成功生成并保存至: {output_path}')
